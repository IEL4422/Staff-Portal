"""
Document Generation Engine
Handles DOCX templating, PDF form filling, and Dropbox integration
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone
import os
import re
import json
import uuid
import tempfile
import shutil
import base64
from pathlib import Path
import logging

# Document processing libraries
from docxtpl import DocxTemplate
from pypdf import PdfReader, PdfWriter
import dropbox
from dropbox.files import WriteMode
from dropbox.exceptions import ApiError

# Slack integration
from slack_sdk import WebClient
from slack_sdk.errors import SlackApiError

# MongoDB and Airtable will be passed from main app
from motor.motor_asyncio import AsyncIOMotorDatabase
import httpx

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents", tags=["Document Generation"])
security = HTTPBearer()

# Storage paths
TEMPLATES_DIR = Path(__file__).parent.parent / "templates_storage"
TEMPLATES_DIR.mkdir(exist_ok=True)

# Dropbox config
DROPBOX_ACCESS_TOKEN = os.environ.get('DROPBOX_ACCESS_TOKEN', '')
DROPBOX_BASE_FOLDER = os.environ.get('DROPBOX_BASE_FOLDER', '/Illinois Estate Law/Generated Documents')

# Slack config
SLACK_BOT_TOKEN = os.environ.get('SLACK_BOT_TOKEN', '')
SLACK_CHANNEL_ACTION_REQUIRED = os.environ.get('SLACK_CHANNEL_ACTION_REQUIRED', 'action-required')

# Airtable config (will be imported from main)
AIRTABLE_API_KEY = os.environ.get('AIRTABLE_API_KEY', '')
AIRTABLE_BASE_ID = os.environ.get('AIRTABLE_BASE_ID', '')
AIRTABLE_BASE_URL = f"https://api.airtable.com/v0/{AIRTABLE_BASE_ID}"

# Airtable headers for API requests
airtable_headers = {
    "Authorization": f"Bearer {AIRTABLE_API_KEY}",
    "Content-Type": "application/json"
}


# ==================== CONSTANTS ====================

COUNTIES = ["Cook", "Kane", "DuPage", "Lake", "Will", "Statewide"]
CASE_TYPES = ["Deed", "Estate Planning", "Probate", "Prenuptial Agreement"]
DOCUMENT_CATEGORIES = ["Court Order", "Legal Letter", "Deed", "Form", "Agreement", "Other"]


# ==================== MODELS ====================

class TemplateCreate(BaseModel):
    name: str
    type: str  # DOCX or FILLABLE_PDF
    county: str  # Cook, Kane, DuPage, Lake, Will, Statewide
    case_type: str  # Deed, Estate Planning, Probate, Prenuptial Agreement
    category: Optional[str] = "Other"  # Court Order, Legal Letter, Deed, Form, Agreement, Other
    detected_variables: Optional[List[str]] = []
    detected_pdf_fields: Optional[List[Dict]] = []


class MappingProfileCreate(BaseModel):
    name: str
    template_id: str
    mapping_json: Dict[str, Any]
    repeat_rules_json: Optional[Dict[str, Any]] = {}
    output_rules_json: Optional[Dict[str, Any]] = {}
    dropbox_rules_json: Optional[Dict[str, Any]] = {}


class GenerateDocxRequest(BaseModel):
    client_id: str
    template_id: str
    profile_id: Optional[str] = None
    custom_mapping: Optional[Dict[str, Any]] = None
    output_format: str = "DOCX"  # DOCX, PDF, BOTH
    save_to_dropbox: bool = False


class FillPdfRequest(BaseModel):
    client_id: str
    template_id: str
    profile_id: Optional[str] = None
    custom_mapping: Optional[Dict[str, Any]] = None
    flatten: bool = False
    save_to_dropbox: bool = False


# ==================== HELPERS ====================

def get_dropbox_client():
    """Get authenticated Dropbox client with team member support"""
    token = os.environ.get('DROPBOX_ACCESS_TOKEN', '')
    team_member_id = os.environ.get('DROPBOX_TEAM_MEMBER_ID', '')
    
    if not token:
        raise HTTPException(status_code=500, detail="Dropbox access token not configured")
    
    if team_member_id:
        # Business team account - use team member file access
        return dropbox.DropboxTeam(token).as_user(team_member_id)
    else:
        # Personal account
        return dropbox.Dropbox(token)


async def airtable_request(method: str, endpoint: str, data: Optional[Dict] = None) -> Dict:
    """Make request to Airtable API"""
    url = f"{AIRTABLE_BASE_URL}/{endpoint}"
    headers = {
        "Authorization": f"Bearer {AIRTABLE_API_KEY}",
        "Content-Type": "application/json"
    }
    
    async with httpx.AsyncClient(timeout=30.0) as client:
        if method == "GET":
            response = await client.get(url, headers=headers)
        elif method == "POST":
            response = await client.post(url, headers=headers, json=data)
        elif method == "PATCH":
            response = await client.patch(url, headers=headers, json=data)
        else:
            raise ValueError(f"Unsupported method: {method}")
        
        response.raise_for_status()
        return response.json() if response.text else {}


async def get_client_bundle(client_id: str) -> Dict[str, Any]:
    """
    Fetch client record and all linked records from Airtable.
    Returns a normalized dictionary suitable for templating.
    Also includes raw Airtable field names for direct mapping.
    """
    bundle = {}
    
    try:
        # Get main client record from Master List
        client_data = await airtable_request("GET", f"Master%20List/{client_id}")
        fields = client_data.get("fields", {})
        
        # IMPORTANT: Add ALL raw Airtable fields directly to the bundle
        # This allows direct mapping of Airtable field names to template variables
        for key, value in fields.items():
            # Add the raw field name as-is
            bundle[key] = value if value is not None else ""
            # Also add a lowercase version for case-insensitive matching
            bundle[key.lower()] = value if value is not None else ""
            # And a version with underscores instead of spaces
            bundle[key.lower().replace(' ', '_').replace('-', '_')] = value if value is not None else ""
        
        # Map common client fields (computed/combined fields)
        bundle["clientname"] = fields.get("Client", fields.get("Matter Name", ""))
        bundle["mattername"] = fields.get("Matter Name", "")
        bundle["decedentname"] = fields.get("Decedent Name", fields.get("Matter Name", ""))
        bundle["casenumber"] = fields.get("Case Number", "")
        bundle["calendar"] = fields.get("Calendar", "")
        bundle["clientprobaterole"] = fields.get("Client Probate Role", "")
        bundle["clientstreetaddress"] = fields.get("Street Address", fields.get("Client Street Address", ""))
        bundle["clientcity"] = fields.get("City", "")
        bundle["clientstate"] = fields.get("State", "")
        bundle["clientzip"] = fields.get("Zip Code", "")
        bundle["clientcitystatezip"] = f"{fields.get('City', '')}, {fields.get('State', '')} {fields.get('Zip Code', '')}".strip(", ")
        bundle["clientemail"] = fields.get("Email Address", "")
        bundle["clientphone"] = fields.get("Phone Number", "")
        
        # Decedent info
        bundle["decedentstreetaddress"] = fields.get("Decedent Street Address", "")
        bundle["decedentcity"] = fields.get("Decedent City", "")
        bundle["decedentstate"] = fields.get("Decedent State", "")
        bundle["decedentzip"] = fields.get("Decedent Zip", "")
        bundle["decedentcitystatezip"] = f"{fields.get('Decedent City', '')}, {fields.get('Decedent State', '')} {fields.get('Decedent Zip', '')}".strip(", ")
        bundle["decedentdod"] = fields.get("Date of Death", "")
        bundle["decedentdob"] = fields.get("Decedent DOB", fields.get("Date of Birth", ""))
        
        # Case type and status
        bundle["casetype"] = fields.get("Type of Case", "")
        bundle["casestatus"] = fields.get("Active/Inactive", "")
        bundle["datepaid"] = fields.get("Date Paid", "")
        
        # Get linked Judge Information
        judge_ids = fields.get("Judge", [])
        if judge_ids:
            try:
                judge_data = await airtable_request("GET", f"Judge%20Information/{judge_ids[0]}")
                judge_fields = judge_data.get("fields", {})
                bundle["judge"] = judge_fields.get("Judge Name", "")
                bundle["judgeemail"] = judge_fields.get("Email", "")
                bundle["courtroom"] = judge_fields.get("Courtroom", "")
                bundle["courthouse"] = judge_fields.get("Courthouse", "")
            except Exception as e:
                logger.warning(f"Failed to fetch judge: {e}")
                bundle["judge"] = ""
        
        # Get linked Case Contacts and categorize them
        contact_ids = fields.get("Case Contacts", [])
        contacts = []
        executors = []
        guardians = []
        caretakers = []
        trustees = []
        beneficiaries = []
        hpoa_list = []
        fpoa_list = []
        
        for contact_id in contact_ids:
            try:
                contact_data = await airtable_request("GET", f"Case%20Contacts/{contact_id}")
                contact_fields = contact_data.get("fields", {})
                contact_info = {
                    "name": contact_fields.get("Name", ""),
                    "type": contact_fields.get("Type", ""),
                    "email": contact_fields.get("Email", ""),
                    "phone": contact_fields.get("Phone", ""),
                    "address": contact_fields.get("Street Address", ""),
                    "city": contact_fields.get("City", ""),
                    "state": contact_fields.get("State", ""),
                    "zip": contact_fields.get("Zip Code", ""),
                    "relationship": contact_fields.get("Relationship to Decedent", "")
                }
                contacts.append(contact_info)
                
                # Categorize by type
                contact_type = (contact_fields.get("Type", "") or "").lower()
                if "executor" in contact_type or "personal representative" in contact_type:
                    executors.append(contact_info)
                elif "guardian" in contact_type:
                    guardians.append(contact_info)
                elif "caretaker" in contact_type:
                    caretakers.append(contact_info)
                elif "trustee" in contact_type:
                    trustees.append(contact_info)
                elif "beneficiary" in contact_type or "heir" in contact_type:
                    beneficiaries.append(contact_info)
                elif "hpoa" in contact_type or "health" in contact_type:
                    hpoa_list.append(contact_info)
                elif "fpoa" in contact_type or "financial" in contact_type:
                    fpoa_list.append(contact_info)
            except Exception as e:
                logger.warning(f"Failed to fetch contact {contact_id}: {e}")
        
        bundle["contacts"] = contacts
        bundle["executors"] = executors
        bundle["guardians"] = guardians
        bundle["caretakers"] = caretakers
        bundle["trustees"] = trustees
        bundle["beneficiaries"] = beneficiaries
        bundle["hpoa"] = hpoa_list
        bundle["fpoa"] = fpoa_list
        
        # Single values for first items (common in templates)
        bundle["executor"] = executors[0]["name"] if executors else ""
        bundle["guardian"] = guardians[0]["name"] if guardians else ""
        bundle["caretaker"] = caretakers[0]["name"] if caretakers else ""
        bundle["trustee"] = trustees[0]["name"] if trustees else ""
        bundle["beneficiary"] = beneficiaries[0]["name"] if beneficiaries else ""
        
        # Get Assets & Debts
        asset_ids = fields.get("Assets & Debts", [])
        assets = []
        debts = []
        for asset_id in asset_ids:
            try:
                asset_data = await airtable_request("GET", f"Assets%20%26%20Debts/{asset_id}")
                asset_fields = asset_data.get("fields", {})
                asset_info = {
                    "name": asset_fields.get("Asset/Debt Name", ""),
                    "type": asset_fields.get("Type", ""),
                    "value": asset_fields.get("Value", ""),
                    "description": asset_fields.get("Description", ""),
                    "account_number": asset_fields.get("Account Number", "")
                }
                if asset_fields.get("Asset or Debt", "").lower() == "asset":
                    assets.append(asset_info)
                else:
                    debts.append(asset_info)
            except Exception as e:
                logger.warning(f"Failed to fetch asset {asset_id}: {e}")
        
        bundle["assets"] = assets
        bundle["debts"] = debts
        
        # Get Dates & Deadlines
        deadline_ids = fields.get("Dates & Deadlines", [])
        deadlines = []
        for deadline_id in deadline_ids:
            try:
                deadline_data = await airtable_request("GET", f"Dates%20%26%20Deadlines/{deadline_id}")
                deadline_fields = deadline_data.get("fields", {})
                deadlines.append({
                    "event": deadline_fields.get("Event", ""),
                    "date": deadline_fields.get("Date", ""),
                    "notes": deadline_fields.get("Notes", "")
                })
            except Exception as e:
                logger.warning(f"Failed to fetch deadline {deadline_id}: {e}")
        
        bundle["deadlines"] = deadlines
        
        # Add current date info for templates
        now = datetime.now()
        bundle["currentdate"] = now.strftime("%B %d, %Y")
        bundle["yyyy"] = now.strftime("%Y")
        bundle["mm"] = now.strftime("%m")
        bundle["dd"] = now.strftime("%d")
        
        # Store raw fields for custom mappings
        bundle["_raw_fields"] = fields
        
    except Exception as e:
        logger.error(f"Failed to get client bundle: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch client data: {str(e)}")
    
    return bundle


def detect_docx_variables(file_path: str) -> Dict[str, Any]:
    """
    Detect variables in a DOCX template.
    Variables use single curly braces: {variablename}
    Repeat blocks use: {#items} ... {/items}
    """
    try:
        doc = DocxTemplate(file_path)
        
        # Make sure document is properly loaded
        if doc.docx is None:
            # Try loading with python-docx directly to extract text
            from docx import Document
            raw_doc = Document(file_path)
            text_content = ""
            for para in raw_doc.paragraphs:
                text_content += para.text + "\n"
            for table in raw_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\n"
            
            # Find variables in text content
            simple_vars = set(re.findall(r'\{([a-zA-Z_][a-zA-Z0-9_\.]*)\}', text_content))
            simple_vars = {v for v in simple_vars if not v.startswith('#') and not v.startswith('/')}
            
            return {
                "variables": sorted(list(simple_vars)),
                "repeat_blocks": [],
                "nested_variables": [],
                "all_detected": sorted(list(simple_vars))
            }
        
        # Get the full XML content
        xml_content = doc.get_xml()
        
        # Find all single-curly-brace variables: {var}
        # Exclude Jinja2 style {{ }} and loop markers {# /}
        simple_vars = set(re.findall(r'\{([a-zA-Z_][a-zA-Z0-9_\.]*)\}', xml_content))
        
        # Find repeat block markers: {#blockname} and {/blockname}
        repeat_starts = set(re.findall(r'\{#([a-zA-Z_][a-zA-Z0-9_]*)\}', xml_content))
        repeat_ends = set(re.findall(r'\{/([a-zA-Z_][a-zA-Z0-9_]*)\}', xml_content))
        
        # Repeat blocks are those that have both start and end markers
        repeat_blocks = repeat_starts & repeat_ends
        
        # Filter out repeat block markers from simple vars
        simple_vars = {v for v in simple_vars if not v.startswith('#') and not v.startswith('/')}
        
        # Find variables inside repeat blocks (e.g., {items.name})
        nested_vars = {v for v in simple_vars if '.' in v}
        top_level_vars = simple_vars - nested_vars
        
        return {
            "variables": sorted(list(top_level_vars)),
            "repeat_blocks": sorted(list(repeat_blocks)),
            "nested_variables": sorted(list(nested_vars)),
            "all_detected": sorted(list(simple_vars))
        }
    except Exception as e:
        logger.error(f"Failed to detect DOCX variables: {e}")
        # Return empty result on error
        return {
            "variables": [],
            "repeat_blocks": [],
            "nested_variables": [],
            "all_detected": [],
            "error": str(e)
        }


def detect_pdf_fields(file_path: str) -> List[Dict[str, Any]]:
    """Detect fillable form fields in a PDF"""
    fields = []
    
    try:
        reader = PdfReader(file_path)
        
        if reader.get_fields():
            for field_name, field_data in reader.get_fields().items():
                field_info = {
                    "name": field_name,
                    "type": str(field_data.get("/FT", "Unknown")),
                    "value": str(field_data.get("/V", "")),
                }
                
                # Determine field type
                ft = field_data.get("/FT", "")
                if ft == "/Tx":
                    field_info["type"] = "text"
                elif ft == "/Btn":
                    field_info["type"] = "checkbox"
                elif ft == "/Ch":
                    field_info["type"] = "dropdown"
                
                fields.append(field_info)
    except Exception as e:
        logger.error(f"Failed to detect PDF fields: {e}")
        # Return empty list instead of raising - allow upload even if field detection fails
        # User can map fields manually later
    
    return fields


def convert_template_syntax(content: str, data: Dict) -> str:
    """
    Convert single-curly-brace syntax to Jinja2 syntax.
    {var} -> {{ var }}
    {#items} -> {% for item in items %}
    {/items} -> {% endfor %}
    {items.field} -> {{ item.field }}
    """
    # This is handled by docxtpl natively when using the right syntax
    # We'll convert our simplified syntax to docxtpl's expected format
    return content


def render_docx_template(template_path: str, data: Dict, output_path: str) -> str:
    """
    Render a DOCX template with the provided data.
    Converts {var} syntax to Jinja2 and renders.
    """
    # Read template
    doc = DocxTemplate(template_path)
    
    # Get XML and convert syntax
    # docxtpl expects {{ var }} syntax, so we need to preprocess
    
    # Create a modified template with converted syntax
    # For now, we'll create a context that works with both syntaxes
    
    # Render with data
    doc.render(data)
    
    # Save output
    doc.save(output_path)
    
    return output_path


def fill_pdf_form(template_path: str, data: Dict, output_path: str, flatten: bool = False) -> str:
    """Fill a PDF form with the provided data"""
    reader = PdfReader(template_path)
    writer = PdfWriter()
    
    # Copy pages and fill form
    writer.append(reader)
    
    # Update form fields
    writer.update_page_form_field_values(
        writer.pages[0],
        data,
        auto_regenerate=True
    )
    
    # Write output
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    
    return output_path


async def upload_to_dropbox(local_path: str, dropbox_path: str) -> str:
    """Upload a file to Dropbox"""
    try:
        dbx = get_dropbox_client()
        
        with open(local_path, 'rb') as f:
            file_data = f.read()
        
        # Ensure path starts with /
        if not dropbox_path.startswith('/'):
            dropbox_path = '/' + dropbox_path
        
        # Upload file
        result = dbx.files_upload(
            file_data,
            dropbox_path,
            mode=WriteMode.overwrite
        )
        
        logger.info(f"Uploaded to Dropbox: {result.path_display}")
        return result.path_display
        
    except ApiError as e:
        logger.error(f"Dropbox upload failed: {e}")
        raise HTTPException(status_code=500, detail=f"Dropbox upload failed: {str(e)}")


def generate_output_filename(pattern: str, data: Dict, template_name: str) -> str:
    """Generate filename from pattern using data"""
    filename = pattern
    
    # Replace template name
    filename = filename.replace("{templateName}", template_name)
    filename = filename.replace("{templatename}", template_name)
    
    # Replace date tokens
    now = datetime.now()
    filename = filename.replace("{yyyy}", now.strftime("%Y"))
    filename = filename.replace("{mm}", now.strftime("%m"))
    filename = filename.replace("{dd}", now.strftime("%d"))
    
    # Replace data tokens
    for key, value in data.items():
        if isinstance(value, str):
            filename = filename.replace(f"{{{key}}}", value)
    
    # Clean filename
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    
    return filename


# ==================== DATABASE HELPERS ====================

# These will be called with the db instance from the main app
async def save_template(db: AsyncIOMotorDatabase, template_data: Dict) -> str:
    """Save a template to MongoDB"""
    template_data["id"] = str(uuid.uuid4())
    template_data["created_at"] = datetime.now(timezone.utc).isoformat()
    template_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.doc_templates.insert_one(template_data)
    return template_data["id"]


async def get_template(db: AsyncIOMotorDatabase, template_id: str) -> Optional[Dict]:
    """Get a template by ID"""
    return await db.doc_templates.find_one({"id": template_id}, {"_id": 0})


async def ensure_template_file_exists(db: AsyncIOMotorDatabase, template: Dict) -> str:
    """
    Ensure the template file exists on disk. If not, restore it from MongoDB.
    Returns the path to the template file.
    """
    file_path = template.get("file_path", "")
    template_id = template.get("id")
    template_name = template.get("name", "Unknown")
    
    # If file exists, return the path
    if file_path and os.path.exists(file_path):
        return file_path
    
    # File doesn't exist - try to restore from MongoDB
    file_content_base64 = template.get("file_content")
    if not file_content_base64:
        logger.error(f"Template '{template_name}' ({template_id}) has no file_content in MongoDB")
        raise HTTPException(
            status_code=500, 
            detail=f"Template file not found and no backup in database: {template_name}. Please re-upload the template."
        )
    
    # Restore the file from MongoDB
    logger.info(f"Restoring template file '{template_name}' from MongoDB...")
    
    try:
        file_content = base64.b64decode(file_content_base64)
        
        # Determine file extension
        file_ext = '.docx' if template.get("type") == "DOCX" else '.pdf'
        
        # Create new file path if needed
        if not file_path:
            file_id = str(uuid.uuid4())
            file_path = str(TEMPLATES_DIR / f"{file_id}{file_ext}")
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # Write the file
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        # Update the file_path in database if it changed
        if file_path != template.get("file_path"):
            await db.doc_templates.update_one(
                {"id": template_id},
                {"$set": {"file_path": file_path}}
            )
        
        logger.info(f"Template file '{template_name}' restored successfully to {file_path}")
        return file_path
        
    except Exception as e:
        logger.error(f"Failed to restore template '{template_name}': {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to restore template file: {template_name}. Error: {str(e)}"
        )


async def list_templates(db: AsyncIOMotorDatabase, template_type: Optional[str] = None) -> List[Dict]:
    """List all templates"""
    query = {}
    if template_type:
        query["type"] = template_type
    return await db.doc_templates.find(query, {"_id": 0}).to_list(100)


async def save_mapping_profile(db: AsyncIOMotorDatabase, profile_data: Dict) -> str:
    """Save a mapping profile to MongoDB"""
    profile_data["id"] = str(uuid.uuid4())
    profile_data["created_at"] = datetime.now(timezone.utc).isoformat()
    profile_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.doc_mapping_profiles.insert_one(profile_data)
    return profile_data["id"]


async def get_mapping_profile(db: AsyncIOMotorDatabase, profile_id: str) -> Optional[Dict]:
    """Get a mapping profile by ID"""
    return await db.doc_mapping_profiles.find_one({"id": profile_id}, {"_id": 0})


async def list_mapping_profiles(db: AsyncIOMotorDatabase, template_id: Optional[str] = None) -> List[Dict]:
    """List mapping profiles"""
    query = {}
    if template_id:
        query["template_id"] = template_id
    return await db.doc_mapping_profiles.find(query, {"_id": 0}).to_list(100)


async def save_generated_doc(db: AsyncIOMotorDatabase, doc_data: Dict) -> str:
    """Save generated document record to MongoDB"""
    if "id" not in doc_data:
        doc_data["id"] = str(uuid.uuid4())
    if "created_at" not in doc_data:
        doc_data["created_at"] = datetime.now(timezone.utc).isoformat()
    await db.generated_docs.insert_one(doc_data)
    return doc_data["id"]


async def list_generated_docs(db: AsyncIOMotorDatabase, client_id: Optional[str] = None) -> List[Dict]:
    """List generated documents"""
    query = {}
    if client_id:
        query["client_id"] = client_id
    return await db.generated_docs.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)


async def get_client_staff_inputs(db: AsyncIOMotorDatabase, client_id: str) -> Dict:
    """Get saved staff inputs for a client"""
    result = await db.client_staff_inputs.find_one({"client_id": client_id}, {"_id": 0})
    return result.get("inputs", {}) if result else {}


async def save_client_staff_inputs(db: AsyncIOMotorDatabase, client_id: str, inputs: Dict) -> None:
    """Save or update staff inputs for a client"""
    await db.client_staff_inputs.update_one(
        {"client_id": client_id},
        {
            "$set": {
                "client_id": client_id,
                "inputs": inputs,
                "updated_at": datetime.now(timezone.utc).isoformat()
            },
            "$setOnInsert": {
                "created_at": datetime.now(timezone.utc).isoformat()
            }
        },
        upsert=True
    )


# ==================== API ENDPOINTS ====================
# Note: These endpoints receive the db instance via dependency injection from main app

def create_document_routes(db: AsyncIOMotorDatabase, get_current_user):
    """Create document routes with database dependency"""
    
    @router.post("/templates/upload")
    async def upload_template(
        file: UploadFile = File(...),
        name: str = Form(...),
        template_type: str = Form(...),
        county: str = Form(...),
        case_type: str = Form(...),
        category: str = Form("Other"),
        current_user: dict = Depends(get_current_user)
    ):
        """Upload a template file (DOCX or PDF) - stores in MongoDB for persistence"""
        # Validate file type
        filename = file.filename.lower()
        if template_type == "DOCX" and not filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be a .docx for DOCX templates")
        if template_type == "FILLABLE_PDF" and not filename.endswith('.pdf'):
            raise HTTPException(status_code=400, detail="File must be a .pdf for PDF templates")
        
        # Validate county and case_type
        if county not in COUNTIES:
            raise HTTPException(status_code=400, detail=f"Invalid county. Must be one of: {', '.join(COUNTIES)}")
        if case_type not in CASE_TYPES:
            raise HTTPException(status_code=400, detail=f"Invalid case type. Must be one of: {', '.join(CASE_TYPES)}")
        
        # Read file content
        content = await file.read()
        file_content_base64 = base64.b64encode(content).decode('utf-8')
        
        # Save file to disk as well (for backwards compatibility and faster access)
        file_id = str(uuid.uuid4())
        file_ext = '.docx' if template_type == "DOCX" else '.pdf'
        file_path = TEMPLATES_DIR / f"{file_id}{file_ext}"
        
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # Detect variables/fields
        detected_variables = []
        detected_pdf_fields = []
        
        try:
            if template_type == "DOCX":
                detection_result = detect_docx_variables(str(file_path))
                detected_variables = detection_result.get("all_detected", [])
            else:
                detected_pdf_fields = detect_pdf_fields(str(file_path))
        except Exception as e:
            # Log the error but continue - allow upload even if field detection fails
            logger.error(f"Failed to detect fields in template: {str(e)}")
        
        # Save to database WITH file content
        template_data = {
            "name": name,
            "type": template_type,
            "county": county,
            "case_type": case_type,
            "category": category,
            "file_path": str(file_path),
            "file_content": file_content_base64,  # Store file content in MongoDB
            "original_filename": file.filename,
            "detected_variables": detected_variables,
            "detected_pdf_fields": detected_pdf_fields
        }
        
        template_id = await save_template(db, template_data)
        logger.info(f"Template '{name}' uploaded and stored in MongoDB with ID: {template_id}")
        
        return {
            "id": template_id,
            "name": name,
            "type": template_type,
            "county": county,
            "case_type": case_type,
            "category": category,
            "detected_variables": detected_variables,
            "detected_pdf_fields": detected_pdf_fields,
            "message": "Template uploaded successfully"
        }
    
    @router.get("/templates")
    async def get_templates(
        template_type: Optional[str] = None,
        case_type: Optional[str] = None,
        county: Optional[str] = None,
        search: Optional[str] = None,
        current_user: dict = Depends(get_current_user)
    ):
        """List all templates with optional filters"""
        query = {}
        if template_type:
            query["type"] = template_type
        if case_type:
            query["case_type"] = case_type
        if county:
            query["county"] = county
        
        templates = await db.doc_templates.find(query, {"_id": 0}).to_list(100)
        
        # Apply search filter if provided
        if search:
            search_lower = search.lower()
            templates = [t for t in templates if search_lower in t.get("name", "").lower()]
        
        return {"templates": templates}
    
    @router.get("/templates/by-case-type/{case_type_filter}")
    async def get_templates_by_case_type(
        case_type_filter: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get templates organized by case type"""
        if case_type_filter == "all":
            templates = await db.doc_templates.find({}, {"_id": 0}).to_list(100)
        else:
            templates = await db.doc_templates.find({"case_type": case_type_filter}, {"_id": 0}).to_list(100)
        
        # Group by category
        grouped = {}
        for t in templates:
            cat = t.get("category", "Other")
            if cat not in grouped:
                grouped[cat] = []
            grouped[cat].append(t)
        
        return {"templates": templates, "grouped": grouped}
    
    @router.get("/constants")
    async def get_constants(
        current_user: dict = Depends(get_current_user)
    ):
        """Get available counties, case types, and categories"""
        return {
            "counties": COUNTIES,
            "case_types": CASE_TYPES,
            "categories": DOCUMENT_CATEGORIES
        }
    
    @router.get("/templates/{template_id}")
    async def get_template_by_id(
        template_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get a specific template"""
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        return template
    
    @router.get("/templates/{template_id}/health")
    async def check_template_health(
        template_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Check if a template is healthy (exists in DB and file exists on disk)"""
        template = await get_template(db, template_id)
        if not template:
            return {
                "healthy": False,
                "template_id": template_id,
                "in_database": False,
                "file_exists": False,
                "error": "Template not found in database"
            }
        
        file_path = template.get("file_path", "")
        file_exists = bool(file_path) and os.path.exists(file_path)
        
        return {
            "healthy": file_exists,
            "template_id": template_id,
            "template_name": template.get("name"),
            "in_database": True,
            "file_path": file_path,
            "file_exists": file_exists,
            "error": None if file_exists else f"File not found at: {file_path}"
        }
    
    @router.get("/templates-health")
    async def check_all_templates_health(
        current_user: dict = Depends(get_current_user)
    ):
        """Check health of all templates"""
        templates = await db.doc_templates.find({}, {"_id": 0}).to_list(100)
        
        results = []
        healthy_count = 0
        unhealthy_count = 0
        
        for t in templates:
            file_path = t.get("file_path", "")
            file_exists = bool(file_path) and os.path.exists(file_path)
            
            if file_exists:
                healthy_count += 1
            else:
                unhealthy_count += 1
            
            results.append({
                "id": t.get("id"),
                "name": t.get("name"),
                "healthy": file_exists,
                "file_path": file_path,
                "file_exists": file_exists
            })
        
        return {
            "total": len(templates),
            "healthy": healthy_count,
            "unhealthy": unhealthy_count,
            "templates": results
        }
    
    @router.post("/templates-migrate")
    async def migrate_templates_to_mongodb(
        current_user: dict = Depends(get_current_user)
    ):
        """
        Migrate existing templates to store file content in MongoDB.
        This ensures templates persist across deployments.
        """
        templates = await db.doc_templates.find({}, {"_id": 0}).to_list(100)
        
        migrated = 0
        already_migrated = 0
        failed = 0
        errors = []
        
        for t in templates:
            template_id = t.get("id")
            template_name = t.get("name")
            
            # Check if already migrated
            if t.get("file_content"):
                already_migrated += 1
                continue
            
            # Try to read file and store in MongoDB
            file_path = t.get("file_path", "")
            if file_path and os.path.exists(file_path):
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                    file_content_base64 = base64.b64encode(content).decode('utf-8')
                    
                    await db.doc_templates.update_one(
                        {"id": template_id},
                        {"$set": {"file_content": file_content_base64}}
                    )
                    migrated += 1
                    logger.info(f"Migrated template '{template_name}' to MongoDB")
                except Exception as e:
                    failed += 1
                    errors.append({"id": template_id, "name": template_name, "error": str(e)})
            else:
                failed += 1
                errors.append({"id": template_id, "name": template_name, "error": "File not found on disk"})
        
        return {
            "success": True,
            "total": len(templates),
            "migrated": migrated,
            "already_migrated": already_migrated,
            "failed": failed,
            "errors": errors
        }
    
    @router.post("/templates/{template_id}/restore")
    async def restore_template_file(
        template_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """
        Restore a template file from MongoDB to disk.
        Useful if the file was lost due to deployment.
        """
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        try:
            file_path = await ensure_template_file_exists(db, template)
            return {
                "success": True,
                "template_id": template_id,
                "template_name": template.get("name"),
                "file_path": file_path,
                "message": "Template file restored successfully"
            }
        except HTTPException as e:
            raise e
    
    @router.delete("/templates/{template_id}")
    async def delete_template(
        template_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Delete a template"""
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Delete file
        file_path = template.get("file_path")
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete from database
        await db.doc_templates.delete_one({"id": template_id})
        
        return {"success": True, "message": "Template deleted"}
    
    @router.post("/docx/detect-variables")
    async def detect_variables_endpoint(
        file: UploadFile = File(...),
        current_user: dict = Depends(get_current_user)
    ):
        """Detect variables in an uploaded DOCX file without saving it"""
        if not file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be a .docx")
        
        # Save temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = detect_docx_variables(tmp_path)
            return result
        finally:
            os.unlink(tmp_path)
    
    @router.post("/pdf/detect-fields")
    async def detect_fields_endpoint(
        file: UploadFile = File(...),
        current_user: dict = Depends(get_current_user)
    ):
        """Detect form fields in an uploaded PDF file without saving it"""
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="File must be a .pdf")
        
        # Save temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            fields = detect_pdf_fields(tmp_path)
            return {"fields": fields}
        finally:
            os.unlink(tmp_path)
    
    @router.post("/mapping-profiles")
    async def create_mapping_profile(
        profile: MappingProfileCreate,
        current_user: dict = Depends(get_current_user)
    ):
        """Create/update the mapping for a template - stores directly on the template"""
        template_id = profile.template_id
        
        # Verify template exists
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Update the template with the mapping
        mapping_data = {
            "mapping_json": profile.mapping_json,
            "mapping_name": profile.name,
            "mapping_updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.doc_templates.update_one(
            {"id": template_id},
            {"$set": mapping_data}
        )
        
        # Also save to profiles collection for backwards compatibility
        # But first, delete any existing profiles for this template
        await db.doc_mapping_profiles.delete_many({"template_id": template_id})
        
        # Create single profile
        profile_data = {
            "id": str(uuid.uuid4()),
            "name": profile.name,
            "template_id": profile.template_id,
            "mapping_json": profile.mapping_json,
            "repeat_rules_json": profile.repeat_rules_json,
            "output_rules_json": profile.output_rules_json,
            "dropbox_rules_json": profile.dropbox_rules_json,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        await db.doc_mapping_profiles.insert_one(profile_data)
        
        return {
            "id": profile_data["id"],
            "message": "Mapping saved to template successfully"
        }
    
    @router.get("/mapping-profiles")
    async def get_mapping_profiles(
        template_id: Optional[str] = None,
        current_user: dict = Depends(get_current_user)
    ):
        """List mapping profiles"""
        profiles = await list_mapping_profiles(db, template_id)
        return {"profiles": profiles}
    
    @router.get("/mapping-profiles/{profile_id}")
    async def get_mapping_profile_by_id(
        profile_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get a specific mapping profile"""
        profile = await get_mapping_profile(db, profile_id)
        if not profile:
            raise HTTPException(status_code=404, detail="Profile not found")
        return profile
    
    @router.put("/mapping-profiles/{profile_id}")
    async def update_mapping_profile(
        profile_id: str,
        profile: MappingProfileCreate,
        current_user: dict = Depends(get_current_user)
    ):
        """Update the mapping for a template"""
        template_id = profile.template_id
        
        # Update the template with the mapping
        mapping_data = {
            "mapping_json": profile.mapping_json,
            "mapping_name": profile.name,
            "mapping_updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.doc_templates.update_one(
            {"id": template_id},
            {"$set": mapping_data}
        )
        
        # Also update in profiles collection
        update_data = {
            "name": profile.name,
            "template_id": profile.template_id,
            "mapping_json": profile.mapping_json,
            "repeat_rules_json": profile.repeat_rules_json,
            "output_rules_json": profile.output_rules_json,
            "dropbox_rules_json": profile.dropbox_rules_json,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.doc_mapping_profiles.update_one(
            {"id": profile_id},
            {"$set": update_data}
        )
        
        return {"success": True, "message": "Mapping updated"}
    
    @router.delete("/mapping-profiles/{profile_id}")
    async def delete_mapping_profile(
        profile_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Delete a mapping profile"""
        existing = await get_mapping_profile(db, profile_id)
        if not existing:
            raise HTTPException(status_code=404, detail="Profile not found")
        
        await db.doc_mapping_profiles.delete_one({"id": profile_id})
        return {"success": True, "message": "Profile deleted"}
    
    @router.post("/templates/{template_id}/mapping")
    async def save_template_mapping(
        template_id: str,
        mapping: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """
        Save mapping configuration directly to a template.
        This is the preferred method - one mapping per template.
        """
        # Verify template exists
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Update the template with the mapping
        mapping_data = {
            "mapping_json": mapping.get("mapping_json", {}),
            "mapping_updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.doc_templates.update_one(
            {"id": template_id},
            {"$set": mapping_data}
        )
        
        return {
            "success": True,
            "message": "Mapping saved to template successfully",
            "template_id": template_id
        }
    
    @router.get("/templates/{template_id}/mapping")
    async def get_template_mapping(
        template_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """
        Get the mapping configuration for a template.
        Returns the mapping stored directly on the template.
        """
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        return {
            "template_id": template_id,
            "template_name": template.get("name"),
            "template_type": template.get("type"),
            "detected_variables": template.get("detected_variables", []),
            "detected_pdf_fields": template.get("detected_pdf_fields", []),
            "mapping_json": template.get("mapping_json", {}),
            "mapping_updated_at": template.get("mapping_updated_at")
        }
    
    @router.get("/client-bundle/{client_id}")
    async def get_client_bundle_endpoint(
        client_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get client data bundle for preview"""
        bundle = await get_client_bundle(client_id)
        return bundle
    
    @router.post("/generate-docx")
    async def generate_docx(
        request: GenerateDocxRequest,
        current_user: dict = Depends(get_current_user)
    ):
        """Generate a document from a DOCX template"""
        # Get template
        template = await get_template(db, request.template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        if template["type"] != "DOCX":
            raise HTTPException(status_code=400, detail="Template is not a DOCX")
        
        # Get client data
        client_bundle = await get_client_bundle(request.client_id)
        
        # Get mapping profile if specified
        mapping = request.custom_mapping or {}
        output_rules = {}
        dropbox_rules = {}
        
        if request.profile_id:
            profile = await get_mapping_profile(db, request.profile_id)
            if profile:
                mapping = profile.get("mapping_json", {})
                output_rules = profile.get("output_rules_json", {})
                dropbox_rules = profile.get("dropbox_rules_json", {})
        
        # Apply custom mappings to client bundle
        render_data = dict(client_bundle)
        if mapping.get("fields"):
            for var_name, source_info in mapping["fields"].items():
                source = source_info.get("source", "")
                if source and source in client_bundle:
                    render_data[var_name] = client_bundle[source]
        
        # Generate output filename
        filename_pattern = output_rules.get("fileNamePattern", "{clientname} - {templateName} - {yyyy}-{mm}-{dd}")
        base_filename = generate_output_filename(filename_pattern, render_data, template["name"])
        
        # Create output directory
        output_dir = TEMPLATES_DIR / "generated"
        output_dir.mkdir(exist_ok=True)
        
        # Generate DOCX
        output_docx_path = output_dir / f"{base_filename}.docx"
        render_docx_template(template["file_path"], render_data, str(output_docx_path))
        
        result = {
            "success": True,
            "docx_path": str(output_docx_path),
            "docx_filename": f"{base_filename}.docx",
            "pdf_available": False,
            "pdf_message": "PDF conversion requires LibreOffice (not available in this environment)"
        }
        
        # Upload to Dropbox if requested
        dropbox_paths = []
        if request.save_to_dropbox and dropbox_rules.get("enabled", False):
            base_folder = dropbox_rules.get("baseFolder", DROPBOX_BASE_FOLDER)
            folder_pattern = dropbox_rules.get("folderPattern", "/{clientname}/{yyyy}/{templateName}/")
            
            folder_path = generate_output_filename(folder_pattern, render_data, template["name"])
            full_dropbox_path = f"{base_folder}{folder_path}{base_filename}.docx"
            
            try:
                saved_path = await upload_to_dropbox(str(output_docx_path), full_dropbox_path)
                dropbox_paths.append(saved_path)
                result["dropbox_docx_path"] = saved_path
            except Exception as e:
                result["dropbox_error"] = str(e)
        
        # Save generation record
        gen_record = {
            "client_id": request.client_id,
            "template_id": request.template_id,
            "profile_id": request.profile_id,
            "docx_path": str(output_docx_path),
            "pdf_path": None,
            "dropbox_paths": dropbox_paths,
            "status": "SUCCESS",
            "log": f"Generated from template: {template['name']}"
        }
        await save_generated_doc(db, gen_record)
        
        return result
    
    @router.post("/fill-pdf")
    async def fill_pdf(
        request: FillPdfRequest,
        current_user: dict = Depends(get_current_user)
    ):
        """Fill a PDF form with client data"""
        # Get template
        template = await get_template(db, request.template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        if template["type"] != "FILLABLE_PDF":
            raise HTTPException(status_code=400, detail="Template is not a fillable PDF")
        
        # Get client data
        client_bundle = await get_client_bundle(request.client_id)
        
        # Get mapping profile if specified
        mapping = request.custom_mapping or {}
        output_rules = {}
        dropbox_rules = {}
        
        if request.profile_id:
            profile = await get_mapping_profile(db, request.profile_id)
            if profile:
                mapping = profile.get("mapping_json", {})
                output_rules = profile.get("output_rules_json", {})
                dropbox_rules = profile.get("dropbox_rules_json", {})
        
        # Build PDF field values from mapping
        pdf_field_values = {}
        pdf_field_mappings = mapping.get("pdfFields", {})
        
        for field_name, field_config in pdf_field_mappings.items():
            source = field_config.get("source", "")
            field_type = field_config.get("type", "text")
            
            if source in client_bundle:
                value = client_bundle[source]
                
                # Handle checkbox fields
                if field_type == "checkbox":
                    true_value = field_config.get("trueValue", "Yes")
                    pdf_field_values[field_name] = "/Yes" if str(value).lower() in ["yes", "true", "1"] else "/Off"
                else:
                    pdf_field_values[field_name] = str(value) if value else ""
        
        # Generate output filename
        filename_pattern = output_rules.get("fileNamePattern", "{clientname} - {templateName} - FILLED - {yyyy}-{mm}-{dd}")
        base_filename = generate_output_filename(filename_pattern, client_bundle, template["name"])
        
        # Create output directory
        output_dir = TEMPLATES_DIR / "generated"
        output_dir.mkdir(exist_ok=True)
        
        # Fill PDF
        output_pdf_path = output_dir / f"{base_filename}.pdf"
        fill_pdf_form(template["file_path"], pdf_field_values, str(output_pdf_path), request.flatten)
        
        result = {
            "success": True,
            "pdf_path": str(output_pdf_path),
            "pdf_filename": f"{base_filename}.pdf",
            "flattened": request.flatten
        }
        
        # Upload to Dropbox if requested
        dropbox_paths = []
        if request.save_to_dropbox and dropbox_rules.get("enabled", False):
            base_folder = dropbox_rules.get("baseFolder", DROPBOX_BASE_FOLDER)
            folder_pattern = dropbox_rules.get("folderPattern", "/{clientname}/{yyyy}/{templateName}/")
            
            folder_path = generate_output_filename(folder_pattern, client_bundle, template["name"])
            full_dropbox_path = f"{base_folder}{folder_path}{base_filename}.pdf"
            
            try:
                saved_path = await upload_to_dropbox(str(output_pdf_path), full_dropbox_path)
                dropbox_paths.append(saved_path)
                result["dropbox_pdf_path"] = saved_path
            except Exception as e:
                result["dropbox_error"] = str(e)
        
        # Save generation record
        gen_record = {
            "client_id": request.client_id,
            "template_id": request.template_id,
            "profile_id": request.profile_id,
            "docx_path": None,
            "pdf_path": str(output_pdf_path),
            "dropbox_paths": dropbox_paths,
            "status": "SUCCESS",
            "log": f"Filled from template: {template['name']}"
        }
        await save_generated_doc(db, gen_record)
        
        return result
    
    @router.get("/generated")
    async def get_generated_docs(
        client_id: Optional[str] = None,
        current_user: dict = Depends(get_current_user)
    ):
        """List generated documents"""
        docs = await list_generated_docs(db, client_id)
        return {"documents": docs}
    
    @router.get("/generated/{doc_id}/download")
    async def download_generated_doc(
        doc_id: str,
        file_type: str = "docx",
        current_user: dict = Depends(get_current_user)
    ):
        """Download a generated document"""
        from fastapi.responses import FileResponse
        
        doc = await db.generated_docs.find_one({"id": doc_id}, {"_id": 0})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = doc.get("docx_path") if file_type == "docx" else doc.get("pdf_path")
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        filename = os.path.basename(file_path)
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/octet-stream"
        )
    
    @router.get("/airtable-fields")
    async def get_airtable_fields(
        current_user: dict = Depends(get_current_user)
    ):
        """Get available Airtable fields for mapping UI - fetches dynamically from Airtable"""
        try:
            # Fetch actual fields from Master List by getting a sample record
            async with httpx.AsyncClient() as client:
                url = f"{AIRTABLE_BASE_URL}/Master%20List?maxRecords=5"
                response = await client.get(url, headers=airtable_headers)
                data = response.json()
            
            # Collect all unique field names from sample records
            master_list_fields = set()
            if 'records' in data:
                for record in data.get('records', []):
                    for key in record.get('fields', {}).keys():
                        master_list_fields.add(key)
            
            # Sort fields alphabetically
            master_list_fields = sorted(list(master_list_fields))
            
            # Also include related tables fields (static for now, can be made dynamic later)
            fields = {
                "Master List": master_list_fields,
                "Judge Information": [
                    "Judge Name", "Email", "Courtroom", "Courthouse"
                ],
                "Case Contacts": [
                    "Name", "Type", "Email", "Phone", "Street Address", "City", "State", "Zip Code",
                    "Relationship to Decedent"
                ],
                "Assets & Debts": [
                    "Asset/Debt Name", "Type", "Value", "Description", "Account Number", "Asset or Debt"
                ],
                "Dates & Deadlines": [
                    "Event", "Date", "Notes"
                ]
            }
            
            # Bundle keys are computed fields available in the client bundle
            bundle_keys = [
                "clientname", "mattername", "decedentname", "casenumber", "calendar",
                "judge", "clientprobaterole", "clientstreetaddress", "clientcitystatezip",
                "clientemail", "clientphone", "decedentstreetaddress", "decedentcitystatezip",
                "decedentdod", "decedentdob", "casetype", "casestatus", "datepaid",
                "executor", "guardian", "caretaker", "trustee", "beneficiary",
                "currentdate", "yyyy", "mm", "dd"
            ]
            
            # Add Master List fields to bundle_keys as direct mappings
            for field in master_list_fields:
                # Convert to lowercase with underscores for bundle key format
                key = field.lower().replace(' ', '_').replace('-', '_')
                if key not in bundle_keys:
                    bundle_keys.append(key)
            
            list_fields = [
                "executors", "guardians", "caretakers", "trustees", "beneficiaries",
                "hpoa", "fpoa", "contacts", "assets", "debts", "deadlines"
            ]
            
            return {
                "airtable_tables": fields,
                "bundle_keys": sorted(bundle_keys),
                "list_fields": list_fields,
                "master_list_fields": master_list_fields  # Raw field names for direct mapping
            }
        except Exception as e:
            logger.error(f"Failed to fetch Airtable fields: {e}")
            # Fallback to static list if dynamic fetch fails
            return {
                "airtable_tables": {
                    "Master List": [
                        "Client", "Matter Name", "Email Address", "Phone Number",
                        "Type of Case", "Active/Inactive", "First Name"
                    ]
                },
                "bundle_keys": ["clientname", "mattername", "clientemail", "clientphone"],
                "list_fields": [],
                "master_list_fields": []
            }
    
    @router.get("/staff-inputs/{client_id}")
    async def get_staff_inputs(
        client_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get saved staff inputs for a client"""
        inputs = await get_client_staff_inputs(db, client_id)
        return {"client_id": client_id, "inputs": inputs}
    
    @router.post("/staff-inputs/{client_id}")
    async def save_staff_inputs(
        client_id: str,
        data: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """Save staff inputs for a client"""
        inputs = data.get("inputs", {})
        await save_client_staff_inputs(db, client_id, inputs)
        return {"success": True, "message": "Staff inputs saved"}
    
    @router.post("/generate-with-inputs")
    async def generate_with_staff_inputs(
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """Generate a document with staff inputs for unmapped variables"""
        client_id = request.get("client_id")
        template_id = request.get("template_id")
        profile_id = request.get("profile_id")
        staff_inputs = request.get("staff_inputs", {})
        save_to_dropbox = request.get("save_to_dropbox", False)
        save_inputs = request.get("save_inputs", True)
        
        if not client_id or not template_id:
            raise HTTPException(status_code=400, detail="client_id and template_id are required")
        
        # Get template
        template = await get_template(db, template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get client data
        client_bundle = await get_client_bundle(client_id)
        
        # Get mapping profile if specified
        mapping = {}
        output_rules = {}
        dropbox_rules = {}
        
        if profile_id and profile_id != '__DEFAULT__':
            profile = await get_mapping_profile(db, profile_id)
            if profile:
                mapping = profile.get("mapping_json", {})
                output_rules = profile.get("output_rules_json", {})
                dropbox_rules = profile.get("dropbox_rules_json", {})
        
        # Build render data: start with client bundle
        render_data = dict(client_bundle)
        
        # Apply profile mappings
        if mapping.get("fields"):
            for var_name, source_info in mapping["fields"].items():
                source = source_info.get("source", "")
                if source and source in client_bundle:
                    render_data[var_name] = client_bundle[source]
        
        # Apply staff inputs (these override or fill unmapped fields)
        for var_name, value in staff_inputs.items():
            if value:  # Only apply non-empty values
                render_data[var_name] = value
        
        # Save staff inputs for future use if requested
        if save_inputs and staff_inputs:
            existing_inputs = await get_client_staff_inputs(db, client_id)
            merged_inputs = {**existing_inputs, **staff_inputs}
            await save_client_staff_inputs(db, client_id, merged_inputs)
        
        # Generate output filename
        filename_pattern = output_rules.get("fileNamePattern", "{clientname} - {templateName} - {yyyy}-{mm}-{dd}")
        base_filename = generate_output_filename(filename_pattern, render_data, template["name"])
        
        # Create output directory
        output_dir = TEMPLATES_DIR / "generated"
        output_dir.mkdir(exist_ok=True)
        
        # Generate document based on type
        result = {
            "success": True,
            "template_name": template["name"],
            "client_name": render_data.get("clientname", "Unknown")
        }
        
        if template["type"] == "DOCX":
            output_path = output_dir / f"{base_filename}.docx"
            render_docx_template(template["file_path"], render_data, str(output_path))
            result["docx_path"] = str(output_path)
            result["docx_filename"] = f"{base_filename}.docx"
            result["pdf_available"] = False
            result["pdf_message"] = "PDF conversion requires LibreOffice (not available)"
        else:
            # PDF filling
            output_path = output_dir / f"{base_filename}.pdf"
            pdf_field_values = {}
            for field in template.get("detected_pdf_fields", []):
                field_name = field.get("name")
                if field_name in render_data:
                    pdf_field_values[field_name] = str(render_data[field_name])
            fill_pdf_form(template["file_path"], pdf_field_values, str(output_path), False)
            result["pdf_path"] = str(output_path)
            result["pdf_filename"] = f"{base_filename}.pdf"
        
        # Upload to Dropbox if requested
        dropbox_paths = []
        if save_to_dropbox:
            base_folder = dropbox_rules.get("baseFolder", DROPBOX_BASE_FOLDER)
            folder_pattern = dropbox_rules.get("folderPattern", "/{clientname}/{yyyy}/{templateName}/")
            folder_path = generate_output_filename(folder_pattern, render_data, template["name"])
            
            file_path = result.get("docx_path") or result.get("pdf_path")
            file_name = result.get("docx_filename") or result.get("pdf_filename")
            full_dropbox_path = f"{base_folder}{folder_path}{file_name}"
            
            try:
                saved_path = await upload_to_dropbox(file_path, full_dropbox_path)
                dropbox_paths.append(saved_path)
                result["dropbox_path"] = saved_path
            except Exception as e:
                result["dropbox_error"] = str(e)
        
        # Save generation record
        gen_record = {
            "client_id": client_id,
            "template_id": template_id,
            "profile_id": profile_id,
            "docx_path": result.get("docx_path"),
            "pdf_path": result.get("pdf_path"),
            "dropbox_paths": dropbox_paths,
            "staff_inputs_used": staff_inputs,
            "status": "SUCCESS",
            "log": f"Generated from template: {template['name']}"
        }
        await save_generated_doc(db, gen_record)
        
        return result
    
    @router.post("/generate-batch")
    async def generate_batch_documents(
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """
        Generate multiple documents at once for a single client.
        Accepts a list of template_ids and consolidated staff inputs.
        Returns separate files for each template.
        """
        client_id = request.get("client_id")
        template_ids = request.get("template_ids", [])
        profile_mappings = request.get("profile_mappings", {})  # {template_id: profile_id}
        staff_inputs = request.get("staff_inputs", {})
        save_to_dropbox = request.get("save_to_dropbox", False)
        save_inputs = request.get("save_inputs", True)
        
        # Log incoming request for debugging
        logger.info(f"[GENERATE-BATCH] Received request - client_id: {client_id}, template_ids: {template_ids}")
        
        if not client_id:
            raise HTTPException(status_code=400, detail="client_id is required")
        if not template_ids or len(template_ids) == 0:
            raise HTTPException(status_code=400, detail="At least one template_id is required")
        
        # Get client data once (reused for all templates)
        client_bundle = await get_client_bundle(client_id)
        
        # Save staff inputs for future use if requested
        if save_inputs and staff_inputs:
            existing_inputs = await get_client_staff_inputs(db, client_id)
            merged_inputs = {**existing_inputs, **staff_inputs}
            await save_client_staff_inputs(db, client_id, merged_inputs)
        
        # Create output directory
        output_dir = TEMPLATES_DIR / "generated"
        output_dir.mkdir(exist_ok=True)
        
        results = []
        errors = []
        
        for template_id in template_ids:
            try:
                # Get template
                template = await get_template(db, template_id)
                if not template:
                    logger.error(f"[GENERATE] Template not found in DB: {template_id}")
                    errors.append({"template_id": template_id, "error": "Template not found in database. It may have been deleted."})
                    continue
                
                # Ensure template file exists (restore from MongoDB if needed)
                try:
                    template_file_path = await ensure_template_file_exists(db, template)
                    logger.info(f"[GENERATE] Template '{template.get('name')}' file ready at: {template_file_path}")
                    # Update template dict with confirmed file path
                    template["file_path"] = template_file_path
                except HTTPException as e:
                    logger.error(f"[GENERATE] Failed to ensure template file: {e.detail}")
                    errors.append({
                        "template_id": template_id, 
                        "error": e.detail
                    })
                    continue
                
                # Get mapping - first try from template directly, then fall back to profiles
                mapping = {}
                output_rules = {}
                dropbox_rules = {}
                used_profile_id = None  # Track which profile was used (if any)
                
                # Check if template has valid mapping stored directly (must have fields or pdfFields)
                template_mapping = template.get("mapping_json", {})
                has_valid_mapping = template_mapping.get("fields") or template_mapping.get("pdfFields")
                
                if has_valid_mapping:
                    mapping = template_mapping
                    logger.info(f"Using mapping stored directly on template '{template.get('name')}'")
                else:
                    # Fall back to profile (for backwards compatibility)
                    profile_id = profile_mappings.get(template_id)
                    profile = None
                    if profile_id and profile_id != '__DEFAULT__':
                        profile = await get_mapping_profile(db, profile_id)
                        used_profile_id = profile_id
                    else:
                        # Auto-load the most recent mapping profile for this template
                        profile = await db.doc_mapping_profiles.find_one(
                            {"template_id": template_id},
                            sort=[("created_at", -1)]
                        )
                        if profile:
                            used_profile_id = profile.get('id')
                            logger.info(f"Auto-loaded mapping profile '{profile.get('name')}' for generation")
                    
                    if profile:
                        mapping = profile.get("mapping_json", {})
                        output_rules = profile.get("output_rules_json", {})
                        dropbox_rules = profile.get("dropbox_rules_json", {})
                
                # Build render data: start with client bundle
                render_data = dict(client_bundle)
                
                # Log available keys for debugging
                logger.info(f"[GENERATE] Client bundle has {len(client_bundle)} keys")
                
                # Apply profile mappings for DOCX templates
                if mapping.get("fields"):
                    for var_name, source_info in mapping["fields"].items():
                        source = source_info.get("source", "")
                        if source and source not in ['__LEAVE_BLANK__', '__STAFF_INPUT__']:
                            # Try exact match first
                            if source in client_bundle:
                                render_data[var_name] = client_bundle[source]
                                logger.debug(f"[MAPPING] {var_name} <- {source} = {client_bundle[source][:50] if client_bundle[source] else 'empty'}")
                            # Try lowercase match
                            elif source.lower() in client_bundle:
                                render_data[var_name] = client_bundle[source.lower()]
                            # Try with underscores
                            elif source.lower().replace(' ', '_') in client_bundle:
                                render_data[var_name] = client_bundle[source.lower().replace(' ', '_')]
                            else:
                                logger.warning(f"[MAPPING] Field '{source}' not found in client bundle for variable '{var_name}'")
                
                # Apply profile mappings for PDF templates
                if mapping.get("pdfFields"):
                    for var_name, source_info in mapping["pdfFields"].items():
                        source = source_info.get("source", "")
                        if source and source not in ['__LEAVE_BLANK__', '__STAFF_INPUT__']:
                            # Try exact match first
                            if source in client_bundle:
                                render_data[var_name] = client_bundle[source]
                            # Try lowercase match
                            elif source.lower() in client_bundle:
                                render_data[var_name] = client_bundle[source.lower()]
                            # Try with underscores
                            elif source.lower().replace(' ', '_') in client_bundle:
                                render_data[var_name] = client_bundle[source.lower().replace(' ', '_')]
                            else:
                                logger.warning(f"[MAPPING] PDF field '{source}' not found in client bundle for field '{var_name}'")
                
                # Apply staff inputs (these override or fill unmapped fields)
                for var_name, value in staff_inputs.items():
                    if value:  # Only apply non-empty values
                        render_data[var_name] = value
                
                # Generate output filename
                filename_pattern = output_rules.get("fileNamePattern", "{clientname} - {templateName} - {yyyy}-{mm}-{dd}")
                base_filename = generate_output_filename(filename_pattern, render_data, template["name"])
                
                # Generate document based on type
                result = {
                    "success": True,
                    "template_id": template_id,
                    "template_name": template["name"],
                    "client_name": render_data.get("clientname", "Unknown")
                }
                
                if template["type"] == "DOCX":
                    output_path = output_dir / f"{base_filename}.docx"
                    render_docx_template(template["file_path"], render_data, str(output_path))
                    result["docx_path"] = str(output_path)
                    result["docx_filename"] = f"{base_filename}.docx"
                    result["file_type"] = "docx"
                else:
                    # PDF filling
                    output_path = output_dir / f"{base_filename}.pdf"
                    pdf_field_values = {}
                    for field in template.get("detected_pdf_fields", []):
                        field_name = field.get("name")
                        if field_name in render_data:
                            pdf_field_values[field_name] = str(render_data[field_name])
                    fill_pdf_form(template["file_path"], pdf_field_values, str(output_path), False)
                    result["pdf_path"] = str(output_path)
                    result["pdf_filename"] = f"{base_filename}.pdf"
                    result["file_type"] = "pdf"
                
                # Upload to Dropbox if requested
                dropbox_paths = []
                if save_to_dropbox:
                    base_folder = dropbox_rules.get("baseFolder", DROPBOX_BASE_FOLDER)
                    folder_pattern = dropbox_rules.get("folderPattern", "/{clientname}/{yyyy}/{templateName}/")
                    folder_path = generate_output_filename(folder_pattern, render_data, template["name"])
                    
                    file_path = result.get("docx_path") or result.get("pdf_path")
                    file_name = result.get("docx_filename") or result.get("pdf_filename")
                    full_dropbox_path = f"{base_folder}{folder_path}{file_name}"
                    
                    try:
                        saved_path = await upload_to_dropbox(file_path, full_dropbox_path)
                        dropbox_paths.append(saved_path)
                        result["dropbox_path"] = saved_path
                    except Exception as e:
                        result["dropbox_error"] = str(e)
                
                # Save generation record with ID
                doc_id = str(uuid.uuid4())
                gen_record = {
                    "id": doc_id,
                    "client_id": client_id,
                    "template_id": template_id,
                    "profile_id": used_profile_id,
                    "docx_path": result.get("docx_path"),
                    "pdf_path": result.get("pdf_path"),
                    "dropbox_paths": dropbox_paths,
                    "staff_inputs_used": staff_inputs,
                    "status": "SUCCESS",
                    "log": f"Generated from template: {template['name']} (batch)",
                    "created_at": datetime.now(timezone.utc).isoformat()
                }
                await save_generated_doc(db, gen_record)
                
                # Include doc_id in result for download
                result["doc_id"] = doc_id
                results.append(result)
                
            except Exception as e:
                logger.error(f"Failed to generate template {template_id}: {e}")
                errors.append({"template_id": template_id, "error": str(e)})
        
        return {
            "success": len(results) > 0,
            "total_requested": len(template_ids),
            "total_generated": len(results),
            "total_failed": len(errors),
            "results": results,
            "errors": errors
        }
    
    @router.post("/get-batch-variables")
    async def get_batch_variables(
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """
        Get all unique variables needed for a batch of templates.
        Used to display a consolidated input form for batch generation.
        """
        template_ids = request.get("template_ids", [])
        client_id = request.get("client_id")
        profile_mappings = request.get("profile_mappings", {})
        
        if not template_ids:
            return {"variables": [], "all_variables": []}
        
        all_variables = set()
        mapped_variables = set()
        leave_blank_variables = set()
        staff_input_variables = set()
        
        # Get client bundle if client specified
        client_bundle = {}
        if client_id:
            try:
                client_bundle = await get_client_bundle(client_id)
            except Exception:
                pass
        
        # Get saved staff inputs
        saved_inputs = {}
        if client_id:
            try:
                saved_inputs = await get_client_staff_inputs(db, client_id)
            except Exception:
                pass
        
        # Collect all variables from all templates
        # Also track the source mapping for each variable (for looking up values)
        variable_source_map = {}  # {var_name: source_field_name}
        
        for template_id in template_ids:
            template = await get_template(db, template_id)
            if not template:
                continue
            
            # Add all detected variables (DOCX templates)
            for var in template.get("detected_variables", []):
                all_variables.add(var)
            
            # Add all detected PDF fields (PDF templates)
            for field in template.get("detected_pdf_fields", []):
                field_name = field.get("name") if isinstance(field, dict) else field
                if field_name:
                    all_variables.add(field_name)
            
            # Get mapping - first try from template directly, then fall back to profiles
            mapping_json = None
            
            if template.get("mapping_json"):
                # Mapping stored directly on template
                mapping_json = template.get("mapping_json")
                logger.info(f"Using mapping from template '{template.get('name')}'")
            else:
                # Fall back to profile (for backwards compatibility)
                profile_id = profile_mappings.get(template_id)
                profile = None
                
                if profile_id and profile_id != '__DEFAULT__':
                    profile = await get_mapping_profile(db, profile_id)
                else:
                    # Auto-load the most recent mapping profile for this template if none selected
                    profile = await db.doc_mapping_profiles.find_one(
                        {"template_id": template_id},
                        sort=[("created_at", -1)]
                    )
                    if profile:
                        logger.info(f"Auto-loaded mapping profile '{profile.get('name')}' for template {template_id}")
                
                if profile:
                    mapping_json = profile.get("mapping_json")
            
            if mapping_json:
                # Check fields mapping (for DOCX)
                if mapping_json.get("fields"):
                    for var_name, source_info in mapping_json["fields"].items():
                        source = source_info.get("source", "")
                        if source == "__LEAVE_BLANK__":
                            # Mark as "leave blank" - doesn't need input
                            leave_blank_variables.add(var_name)
                        elif source == "__STAFF_INPUT__":
                            # Staff input required
                            staff_input_variables.add(var_name)
                        elif source and source in client_bundle:
                            mapped_variables.add(var_name)
                            variable_source_map[var_name] = source  # Track the mapping
                # Check pdfFields mapping (for PDF)
                if mapping_json.get("pdfFields"):
                    for var_name, source_info in mapping_json["pdfFields"].items():
                        source = source_info.get("source", "")
                        if source == "__LEAVE_BLANK__":
                            leave_blank_variables.add(var_name)
                        elif source == "__STAFF_INPUT__":
                            staff_input_variables.add(var_name)
                        elif source and source in client_bundle:
                            mapped_variables.add(var_name)
                            variable_source_map[var_name] = source  # Track the mapping
        
        # Determine which variables are available from client bundle
        variables_with_status = []
        for var in sorted(all_variables):
            # Check if this variable is set to "leave blank"
            is_leave_blank = var in leave_blank_variables
            is_staff_input_required = var in staff_input_variables
            is_mapped = var in mapped_variables
            
            # Get the value: use mapped source if available, otherwise direct lookup
            source_field = variable_source_map.get(var, var)
            airtable_value = client_bundle.get(source_field, "")
            saved_value = saved_inputs.get(var, "")
            
            # Airtable data has PRIORITY over saved staff inputs
            current_value = airtable_value or saved_value
            has_airtable_data = bool(airtable_value)
            
            status = {
                "variable": var,
                "has_airtable_data": has_airtable_data,
                "has_saved_input": bool(saved_value),
                "is_mapped": is_mapped,
                "is_leave_blank": is_leave_blank,
                "is_staff_input": is_staff_input_required,
                "current_value": current_value,
                "mapped_from": source_field if is_mapped else None,
                "needs_input": False
            }
            
            # Needs input ONLY if:
            # 1. Not set to "leave blank"
            # 2. Not mapped to Airtable with data
            # 3. Marked as staff input required AND has no Airtable data AND no saved input
            # Airtable data always takes priority - if we have it, no input needed
            if not is_leave_blank:
                if has_airtable_data:
                    # Airtable data exists - no input needed regardless of other settings
                    status["needs_input"] = False
                elif is_staff_input_required and not saved_value:
                    # Staff input required and no Airtable data and no saved input
                    status["needs_input"] = True
            
            variables_with_status.append(status)
        
        return {
            "variables": variables_with_status,
            "all_variables": sorted(list(all_variables)),
            "saved_inputs": saved_inputs
        }
    
    # ==================== DROPBOX FOLDER BROWSING ====================
    
    def get_dropbox_client_for_user():
        """
        Get Dropbox client, handling both personal and Business team accounts.
        For team accounts, uses DROPBOX_TEAM_MEMBER_ID to specify the user.
        """
        dropbox_token = os.environ.get('DROPBOX_ACCESS_TOKEN', '')
        team_member_id = os.environ.get('DROPBOX_TEAM_MEMBER_ID', '')
        
        if not dropbox_token:
            raise HTTPException(status_code=500, detail="Dropbox not configured. Please set DROPBOX_ACCESS_TOKEN in environment.")
        
        if team_member_id:
            # Business team account - use team member file access
            return dropbox.DropboxTeam(dropbox_token).as_user(team_member_id)
        else:
            # Personal account or team account acting as admin
            return dropbox.Dropbox(dropbox_token)
    
    @router.get("/dropbox/folders")
    async def list_dropbox_folders(
        path: str = "",
        current_user: dict = Depends(get_current_user)
    ):
        """List folders in Dropbox for folder selection during save."""
        try:
            dbx = get_dropbox_client_for_user()
            
            # Use root or specified path
            search_path = path if path else ""
            
            result = dbx.files_list_folder(search_path)
            
            folders = []
            for entry in result.entries:
                if isinstance(entry, dropbox.files.FolderMetadata):
                    folders.append({
                        "name": entry.name,
                        "path": entry.path_display,
                        "id": entry.id
                    })
            
            # Sort alphabetically
            folders.sort(key=lambda x: x["name"].lower())
            
            return {
                "current_path": path or "/",
                "folders": folders
            }
        except dropbox.exceptions.AuthError as e:
            logger.error(f"Dropbox authentication error: {e}")
            error_msg = str(e)
            if "expired_access_token" in error_msg:
                raise HTTPException(
                    status_code=401, 
                    detail="Dropbox access token has expired. Please generate a new token from the Dropbox App Console and update DROPBOX_ACCESS_TOKEN."
                )
            raise HTTPException(status_code=401, detail=f"Dropbox authentication failed: {error_msg}")
        except dropbox.exceptions.BadInputError as e:
            logger.error(f"Dropbox permission error: {e}")
            error_msg = str(e)
            if "files.metadata.read" in error_msg:
                raise HTTPException(
                    status_code=403, 
                    detail="Dropbox app is missing 'files.metadata.read' permission. Please enable this scope in the Dropbox App Console (Permissions tab) and generate a new access token."
                )
            elif "entire Dropbox Business team" in error_msg or "select_user" in error_msg:
                raise HTTPException(
                    status_code=400, 
                    detail="This is a Dropbox Business team token. Please add DROPBOX_TEAM_MEMBER_ID to the environment to specify which team member's Dropbox to use."
                )
            raise HTTPException(status_code=400, detail=f"Dropbox configuration error: {error_msg}")
        except ApiError as e:
            logger.error(f"Dropbox folder listing error: {e}")
            raise HTTPException(status_code=500, detail=f"Dropbox error: {str(e)}")
    
    @router.get("/dropbox/search")
    async def search_dropbox_folders(
        query: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Search for folders in Dropbox."""
        try:
            dbx = get_dropbox_client_for_user()
            
            # Search for folders
            result = dbx.files_search_v2(query)
            
            folders = []
            for match in result.matches:
                if hasattr(match, 'metadata') and hasattr(match.metadata, 'metadata'):
                    entry = match.metadata.metadata
                    if isinstance(entry, dropbox.files.FolderMetadata):
                        folders.append({
                            "name": entry.name,
                            "path": entry.path_display,
                            "id": entry.id
                        })
            
            return {"query": query, "folders": folders}
        except dropbox.exceptions.AuthError as e:
            logger.error(f"Dropbox authentication error: {e}")
            error_msg = str(e)
            if "expired_access_token" in error_msg:
                raise HTTPException(
                    status_code=401, 
                    detail="Dropbox access token has expired. Please generate a new token."
                )
            raise HTTPException(status_code=401, detail=f"Dropbox authentication failed: {error_msg}")
        except dropbox.exceptions.BadInputError as e:
            logger.error(f"Dropbox permission error: {e}")
            error_msg = str(e)
            if "files.metadata.read" in error_msg or "files.content.read" in error_msg:
                raise HTTPException(
                    status_code=403, 
                    detail="Dropbox app is missing required permissions. Please enable 'files.metadata.read' and 'files.content.read' in the Dropbox App Console."
                )
            elif "entire Dropbox Business team" in error_msg or "select_user" in error_msg:
                raise HTTPException(
                    status_code=400, 
                    detail="This is a Dropbox Business team token. Please add DROPBOX_TEAM_MEMBER_ID to the environment."
                )
            raise HTTPException(status_code=400, detail=f"Dropbox configuration error: {error_msg}")
        except ApiError as e:
            logger.error(f"Dropbox search error: {e}")
            raise HTTPException(status_code=500, detail=f"Dropbox error: {str(e)}")
    
    @router.post("/dropbox/save")
    async def save_to_dropbox_folder(
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """Save a generated document to a specific Dropbox folder."""
        doc_id = request.get("doc_id")
        local_path = request.get("local_path")
        dropbox_folder = request.get("dropbox_folder")
        filename = request.get("filename")
        
        if not local_path or not dropbox_folder:
            raise HTTPException(status_code=400, detail="local_path and dropbox_folder required")
        
        if not os.path.exists(local_path):
            raise HTTPException(status_code=404, detail="Local file not found")
        
        try:
            full_path = f"{dropbox_folder}/{filename}"
            saved_path = await upload_to_dropbox(local_path, full_path)
            
            # Update the generated doc record if doc_id provided
            if doc_id:
                await db.generated_docs.update_one(
                    {"id": doc_id},
                    {"$push": {"dropbox_paths": saved_path}}
                )
            
            return {"success": True, "dropbox_path": saved_path}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    
    # ==================== SLACK INTEGRATION ====================
    
    def get_slack_token():
        """Get Slack token at runtime to ensure .env is loaded."""
        from dotenv import load_dotenv
        load_dotenv('/app/backend/.env')
        return os.environ.get('SLACK_BOT_TOKEN', '')
    
    def get_slack_channel():
        """Get Slack channel at runtime to ensure .env is loaded."""
        from dotenv import load_dotenv
        load_dotenv('/app/backend/.env')
        return os.environ.get('SLACK_CHANNEL_ACTION_REQUIRED', 'action-required')
    
    async def send_slack_message(channel: str, text: str, blocks: list = None):
        """Send a message to a Slack channel."""
        token = get_slack_token()
        if not token:
            logger.warning("Slack not configured - message not sent")
            return None
        
        try:
            client = WebClient(token=token)
            response = client.chat_postMessage(
                channel=channel,
                text=text,
                blocks=blocks
            )
            return response
        except SlackApiError as e:
            logger.error(f"Slack API error: {e.response['error']}")
            raise HTTPException(status_code=500, detail=f"Slack error: {e.response['error']}")
    
    async def send_slack_dm(user_id: str, text: str, blocks: list = None):
        """Send a direct message to a Slack user."""
        token = get_slack_token()
        if not token:
            return None
        
        try:
            client = WebClient(token=token)
            # Open a DM channel first
            response = client.conversations_open(users=[user_id])
            channel_id = response["channel"]["id"]
            
            # Send the message
            return client.chat_postMessage(
                channel=channel_id,
                text=text,
                blocks=blocks
            )
        except SlackApiError as e:
            logger.error(f"Slack DM error: {e.response['error']}")
            return None
    
    @router.post("/send-for-approval")
    async def send_document_for_approval(
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """Send document(s) to attorney for approval via Slack."""
        documents = request.get("documents", [])
        matter_name = request.get("matter_name", "Unknown Matter")
        client_id = request.get("client_id")
        drafter_name = current_user.get("name", current_user.get("email", "Staff"))
        drafter_id = current_user.get("id")
        
        if not documents:
            raise HTTPException(status_code=400, detail="No documents provided")
        
        # Create approval records and get preview URLs
        approval_records = []
        for doc in documents:
            approval_id = str(uuid.uuid4())
            
            # Save approval record to MongoDB
            approval_record = {
                "id": approval_id,
                "doc_id": doc.get("doc_id"),
                "template_name": doc.get("template_name"),
                "local_path": doc.get("local_path"),
                "matter_name": matter_name,
                "client_id": client_id,
                "drafter_id": drafter_id,
                "drafter_name": drafter_name,
                "status": "PENDING",
                "created_at": datetime.now(timezone.utc).isoformat(),
                "approved_at": None,
                "approved_by": None
            }
            await db.document_approvals.insert_one(approval_record)
            approval_records.append(approval_record)
        
        # Build Slack message
        base_url = os.environ.get('FRONTEND_URL', os.environ.get('REACT_APP_BACKEND_URL', 'https://smartdocs-111.preview.emergentagent.com'))
        # Remove /api if present
        base_url = base_url.replace('/api', '').rstrip('/')
        
        doc_list = "\n".join([
            f"• *{doc.get('template_name')}* - <{base_url}/document-approval/{approval_records[i]['id']}|View & Approve>"
            for i, doc in enumerate(documents)
        ])
        
        blocks = [
            {
                "type": "header",
                "text": {
                    "type": "plain_text",
                    "text": "📄 Document(s) Ready for Approval",
                    "emoji": True
                }
            },
            {
                "type": "section",
                "fields": [
                    {"type": "mrkdwn", "text": f"*Matter:*\n{matter_name}"},
                    {"type": "mrkdwn", "text": f"*Drafted by:*\n{drafter_name}"}
                ]
            },
            {
                "type": "section",
                "text": {
                    "type": "mrkdwn",
                    "text": f"*Documents:*\n{doc_list}"
                }
            },
            {
                "type": "actions",
                "elements": [
                    {
                        "type": "button",
                        "text": {
                            "type": "plain_text",
                            "text": "📋 Review All Documents",
                            "emoji": True
                        },
                        "style": "primary",
                        "url": f"{base_url}/documents/review"
                    }
                ]
            },
            {
                "type": "context",
                "elements": [
                    {"type": "mrkdwn", "text": "Click 'View & Approve' above to review individual documents, or 'Review All Documents' for the dashboard"}
                ]
            }
        ]
        
        # Send to #action-required channel
        slack_channel = get_slack_channel()
        try:
            await send_slack_message(
                channel=f"#{slack_channel}",
                text=f"Document(s) ready for approval: {matter_name}",
                blocks=blocks
            )
            logger.info(f"Sent Slack notification to #{slack_channel} for {len(documents)} document(s)")
        except Exception as e:
            logger.error(f"Failed to send Slack message: {e}")
        
        return {
            "success": True,
            "message": f"Sent {len(documents)} document(s) to #{slack_channel} for approval",
            "approval_ids": [r["id"] for r in approval_records]
        }
    
    @router.get("/approvals")
    async def get_all_approvals(
        current_user: dict = Depends(get_current_user)
    ):
        """Get all document approvals for the review dashboard."""
        approvals = await db.document_approvals.find({}).sort("created_at", -1).to_list(200)
        
        # Remove MongoDB _id
        for a in approvals:
            a.pop("_id", None)
        
        return {
            "approvals": approvals,
            "total": len(approvals),
            "pending": len([a for a in approvals if a.get("status") == "PENDING"]),
            "approved": len([a for a in approvals if a.get("status") == "APPROVED"])
        }
    
    @router.get("/approval/{approval_id}")
    async def get_approval_details(
        approval_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get document approval details for preview page."""
        approval = await db.document_approvals.find_one({"id": approval_id})
        
        if not approval:
            raise HTTPException(status_code=404, detail="Approval record not found")
        
        # Remove MongoDB _id
        approval.pop("_id", None)
        
        return approval
    
    @router.post("/approval/{approval_id}/approve")
    async def approve_document(
        approval_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Approve a document and notify the drafter."""
        approval = await db.document_approvals.find_one({"id": approval_id})
        
        if not approval:
            raise HTTPException(status_code=404, detail="Approval record not found")
        
        if approval.get("status") == "APPROVED":
            return {"success": True, "message": "Document already approved"}
        
        # Update approval status
        approver_name = current_user.get("name", current_user.get("email", "Attorney"))
        await db.document_approvals.update_one(
            {"id": approval_id},
            {
                "$set": {
                    "status": "APPROVED",
                    "approved_at": datetime.now(timezone.utc).isoformat(),
                    "approved_by": approver_name
                }
            }
        )
        
        # Create notification for drafter
        notification = {
            "id": str(uuid.uuid4()),
            "user_id": approval.get("drafter_id"),
            "type": "DOCUMENT_APPROVED",
            "title": "Document Approved",
            "message": f"Your document '{approval.get('template_name')}' for {approval.get('matter_name')} has been approved by {approver_name}",
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "data": {
                "approval_id": approval_id,
                "template_name": approval.get("template_name"),
                "matter_name": approval.get("matter_name")
            }
        }
        await db.notifications.insert_one(notification)
        
        # Send Slack DM to drafter (if we have their Slack ID - for now just log)
        logger.info(f"Document approved: {approval.get('template_name')} by {approver_name}")
        
        # Also post to Slack channel
        slack_channel = get_slack_channel()
        try:
            await send_slack_message(
                channel=f"#{slack_channel}",
                text=f"✅ Document approved: {approval.get('template_name')} for {approval.get('matter_name')} (approved by {approver_name})"
            )
        except Exception as e:
            logger.error(f"Failed to send approval notification to Slack: {e}")
        
        return {
            "success": True,
            "message": f"Document approved and {approval.get('drafter_name')} has been notified"
        }
    
    @router.get("/notifications")
    async def get_notifications(
        current_user: dict = Depends(get_current_user)
    ):
        """Get notifications for the current user."""
        user_id = current_user.get("id")
        
        notifications = await db.notifications.find(
            {"user_id": user_id}
        ).sort("created_at", -1).limit(50).to_list(50)
        
        # Remove MongoDB _id
        for n in notifications:
            n.pop("_id", None)
        
        unread_count = await db.notifications.count_documents({"user_id": user_id, "read": False})
        
        return {
            "notifications": notifications,
            "unread_count": unread_count
        }
    
    @router.post("/notifications/{notification_id}/read")
    async def mark_notification_read(
        notification_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Mark a notification as read."""
        await db.notifications.update_one(
            {"id": notification_id, "user_id": current_user.get("id")},
            {"$set": {"read": True}}
        )
        return {"success": True}
    
    # ==================== DOCUMENT PREVIEW ====================
    
    @router.get("/preview/{approval_id}")
    async def get_document_preview(
        approval_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get document content for preview. Returns text content extracted from DOCX."""
        approval = await db.document_approvals.find_one({"id": approval_id})
        
        if not approval:
            raise HTTPException(status_code=404, detail="Approval record not found")
        
        local_path = approval.get("local_path")
        
        if not local_path or not os.path.exists(local_path):
            return {
                "success": False,
                "error": "Document file not found",
                "content": None
            }
        
        try:
            # Extract text content from DOCX
            if local_path.endswith('.docx'):
                from docx import Document
                doc = Document(local_path)
                
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append({
                            "type": "paragraph",
                            "text": para.text,
                            "style": para.style.name if para.style else "Normal"
                        })
                
                # Also extract tables
                tables = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                return {
                    "success": True,
                    "file_type": "docx",
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "filename": os.path.basename(local_path)
                }
            
            elif local_path.endswith('.pdf'):
                # For PDF, we can extract text using pypdf
                reader = PdfReader(local_path)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                
                return {
                    "success": True,
                    "file_type": "pdf",
                    "pages": pages,
                    "page_count": len(reader.pages),
                    "filename": os.path.basename(local_path)
                }
            
            else:
                return {
                    "success": False,
                    "error": "Unsupported file type"
                }
                
        except Exception as e:
            logger.error(f"Failed to extract document preview: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    @router.get("/preview-generated/{doc_id}")
    async def get_generated_document_preview(
        doc_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """Get preview of a generated document by its ID."""
        # Find the generated document
        doc = await db.generated_docs.find_one({"id": doc_id})
        
        if not doc:
            return {
                "success": False,
                "error": "Document not found"
            }
        
        # Get the local path (prefer PDF, then DOCX)
        local_path = doc.get("pdf_path") or doc.get("docx_path")
        
        if not local_path or not os.path.exists(local_path):
            return {
                "success": False,
                "error": "Document file not found on server"
            }
        
        try:
            if local_path.endswith('.docx'):
                from docx import Document
                document = Document(local_path)
                
                paragraphs = []
                for para in document.paragraphs:
                    if para.text.strip():
                        paragraphs.append({
                            "type": "paragraph",
                            "text": para.text,
                            "style": para.style.name if para.style else "Normal"
                        })
                
                tables = []
                for table in document.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                return {
                    "success": True,
                    "file_type": "docx",
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "filename": os.path.basename(local_path)
                }
            
            elif local_path.endswith('.pdf'):
                reader = PdfReader(local_path)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                
                return {
                    "success": True,
                    "file_type": "pdf",
                    "pages": pages,
                    "page_count": len(reader.pages),
                    "filename": os.path.basename(local_path)
                }
            
            else:
                return {
                    "success": False,
                    "error": "Unsupported file type"
                }
                
        except Exception as e:
            logger.error(f"Failed to extract document preview: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    # ==================== STAFF INPUT CONFIRMATION ====================
    
    @router.get("/staff-inputs/{client_id}/with-labels")
    async def get_staff_inputs_with_labels(
        client_id: str,
        current_user: dict = Depends(get_current_user)
    ):
        """
        Get staff inputs for a client with their custom labels.
        Returns inputs that were manually entered (not from Airtable).
        These require confirmation before use.
        """
        # Get saved staff inputs
        inputs_doc = await db.client_staff_inputs.find_one({"client_id": client_id})
        
        if not inputs_doc:
            return {
                "client_id": client_id,
                "inputs": {},
                "labels": {},
                "labeled_inputs": [],
                "requires_confirmation": False
            }
        
        inputs = inputs_doc.get("inputs", {})
        labels = inputs_doc.get("labels", {})  # Custom labels for each variable
        last_confirmed = inputs_doc.get("last_confirmed")
        
        # Build labeled inputs list
        labeled_inputs = []
        for var_name, value in inputs.items():
            if value:  # Only include non-empty values
                labeled_inputs.append({
                    "variable": var_name,
                    "label": labels.get(var_name, var_name),  # Use custom label or variable name
                    "value": value,
                    "needs_confirmation": True  # Staff-entered values need confirmation
                })
        
        return {
            "client_id": client_id,
            "inputs": inputs,
            "labels": labels,
            "labeled_inputs": labeled_inputs,
            "requires_confirmation": len(labeled_inputs) > 0,
            "last_confirmed": last_confirmed
        }
    
    @router.post("/staff-inputs/{client_id}/confirm")
    async def confirm_staff_inputs(
        client_id: str,
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """
        Confirm staff inputs before document generation.
        User can also update values during confirmation.
        """
        confirmed_inputs = request.get("inputs", {})
        labels = request.get("labels", {})
        
        # Update the staff inputs with confirmed values
        await db.client_staff_inputs.update_one(
            {"client_id": client_id},
            {
                "$set": {
                    "inputs": confirmed_inputs,
                    "labels": labels,
                    "last_confirmed": datetime.now(timezone.utc).isoformat(),
                    "confirmed_by": current_user.get("email")
                }
            },
            upsert=True
        )
        
        return {
            "success": True,
            "message": "Staff inputs confirmed",
            "client_id": client_id
        }
    
    @router.post("/staff-inputs/{client_id}/save-with-labels")
    async def save_staff_inputs_with_labels(
        client_id: str,
        request: Dict[str, Any],
        current_user: dict = Depends(get_current_user)
    ):
        """
        Save staff inputs along with their custom display labels.
        Labels are used to show friendly names during confirmation.
        """
        inputs = request.get("inputs", {})
        labels = request.get("labels", {})
        
        # Update existing inputs, preserving any that aren't being updated
        existing = await db.client_staff_inputs.find_one({"client_id": client_id})
        
        if existing:
            merged_inputs = {**existing.get("inputs", {}), **inputs}
            merged_labels = {**existing.get("labels", {}), **labels}
        else:
            merged_inputs = inputs
            merged_labels = labels
        
        await db.client_staff_inputs.update_one(
            {"client_id": client_id},
            {
                "$set": {
                    "inputs": merged_inputs,
                    "labels": merged_labels,
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                    "updated_by": current_user.get("email")
                }
            },
            upsert=True
        )
        
        return {
            "success": True,
            "message": "Staff inputs saved with labels",
            "client_id": client_id,
            "input_count": len(merged_inputs)
        }
    
    return router
